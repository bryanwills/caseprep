"""
Export service for CasePrep backend.
Handles exporting transcripts in various formats.
"""

import logging
import tempfile
import os
from pathlib import Path
from typing import Dict, List, Optional, Any
from datetime import datetime
import json

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT

from app.models.transcript import Transcript, TranscriptSegment
from app.models.matter import Matter
from app.core.config import settings

logger = logging.getLogger(__name__)

class ExportService:
    """Service for exporting transcripts in various formats."""

    def __init__(self):
        self.supported_formats = ['srt', 'vtt', 'docx', 'pdf', 'csv', 'json']

    async def export_transcript(
        self,
        transcript: Transcript,
        format: str,
        options: Dict[str, Any],
        matter: Optional[Matter] = None
    ) -> str:
        """
        Export transcript in the specified format.

        Args:
            transcript: The transcript to export
            format: Export format (srt, vtt, docx, pdf, csv, json)
            options: Export options
            matter: Associated matter for additional context

        Returns:
            Path to the exported file
        """
        try:
            if format not in self.supported_formats:
                raise ValueError(f"Unsupported format: {format}")

            logger.info(f"Exporting transcript {transcript.id} in {format} format")

            if format == 'srt':
                return await self._export_srt(transcript, options)
            elif format == 'vtt':
                return await self._export_vtt(transcript, options)
            elif format == 'docx':
                return await self._export_docx(transcript, options, matter)
            elif format == 'pdf':
                return await self._export_pdf(transcript, options, matter)
            elif format == 'csv':
                return await self._export_csv(transcript, options)
            elif format == 'json':
                return await self._export_json(transcript, options)
            else:
                raise ValueError(f"Format {format} not implemented")

        except Exception as e:
            logger.error(f"Export failed for format {format}: {e}")
            raise

    async def _export_srt(self, transcript: Transcript, options: Dict[str, Any]) -> str:
        """Export transcript in SRT (SubRip) format."""
        try:
            output_path = f"/tmp/transcript_{transcript.id}.srt"

            with open(output_path, 'w', encoding='utf-8') as f:
                for i, segment in enumerate(transcript.segments, 1):
                    # SRT format: sequence number, timestamps, text, blank line
                    start_time = self._format_srt_time(segment.start_ms)
                    end_time = self._format_srt_time(segment.end_ms)

                    f.write(f"{i}\n")
                    f.write(f"{start_time} --> {end_time}\n")

                    # Add speaker if enabled
                    if options.get('include_speakers', True):
                        f.write(f"[{segment.speaker}] {segment.text}\n")
                    else:
                        f.write(f"{segment.text}\n")

                    f.write("\n")

            logger.info(f"SRT export completed: {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"SRT export failed: {e}")
            raise

    async def _export_vtt(self, transcript: Transcript, options: Dict[str, Any]) -> str:
        """Export transcript in WebVTT format."""
        try:
            output_path = f"/tmp/transcript_{transcript.id}.vtt"

            with open(output_path, 'w', encoding='utf-8') as f:
                # VTT header
                f.write("WEBVTT\n\n")

                for i, segment in enumerate(transcript.segments, 1):
                    start_time = self._format_vtt_time(segment.start_ms)
                    end_time = self._format_vtt_time(segment.end_ms)

                    f.write(f"{start_time} --> {end_time}\n")

                    # Add speaker if enabled
                    if options.get('include_speakers', True):
                        f.write(f"<v {segment.speaker}>{segment.text}</v>\n")
                    else:
                        f.write(f"{segment.text}\n")

                    f.write("\n")

            logger.info(f"VTT export completed: {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"VTT export failed: {e}")
            raise

    async def _export_docx(self, transcript: Transcript, options: Dict[str, Any], matter: Optional[Matter] = None) -> str:
        """Export transcript in Microsoft Word format."""
        try:
            output_path = f"/tmp/transcript_{transcript.id}.docx"

            # Create document
            doc = Document()

            # Title
            title = doc.add_heading(transcript.title or 'Transcript', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Metadata section
            if options.get('include_metadata', True):
                doc.add_heading('Metadata', level=1)

                metadata_table = doc.add_table(rows=1, cols=2)
                metadata_table.style = 'Table Grid'

                # Add metadata rows
                metadata = [
                    ('Language', transcript.language.upper()),
                    ('Duration', f"{transcript.totalDurationMs / 1000 / 60:.1f} minutes"),
                    ('Model', transcript.asrModel),
                    ('Created', transcript.createdAt),
                ]

                if matter:
                    metadata.extend([
                        ('Case', matter.title),
                        ('Client', matter.clientName or 'N/A'),
                        ('Practice Area', matter.practiceArea or 'N/A'),
                    ])

                for key, value in metadata:
                    row = metadata_table.add_row().cells
                    row[0].text = key
                    row[1].text = str(value)

                doc.add_paragraph()

            # Transcript content
            doc.add_heading('Transcript', level=1)

            for segment in transcript.segments:
                # Speaker and timestamp
                speaker_para = doc.add_paragraph()
                speaker_para.add_run(f"{segment.speaker} ").bold = True
                speaker_para.add_run(f"({self._format_readable_time(segment.start_ms)})").italic = True

                # Text content
                text_para = doc.add_paragraph(segment.text)
                text_para.space_after = Pt(12)

                # Add confidence score if enabled
                if options.get('include_confidence', True) and segment.confidence:
                    confidence_para = doc.add_paragraph()
                    confidence_para.add_run(f"Confidence: {segment.confidence:.2%}").font.size = Pt(8)
                    confidence_para.space_after = Pt(6)

            # Footer with export info
            footer = doc.sections[0].footer
            footer_para = footer.paragraphs[0]
            footer_para.text = f"Exported from CasePrep on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.save(output_path)
            logger.info(f"DOCX export completed: {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"DOCX export failed: {e}")
            raise

    async def _export_pdf(self, transcript: Transcript, options: Dict[str, Any], matter: Optional[Matter] = None) -> str:
        """Export transcript in PDF format (Quote Pack style)."""
        try:
            output_path = f"/tmp/transcript_{transcript.id}.pdf"

            # Create PDF document
            doc = SimpleDocTemplate(
                output_path,
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )

            # Styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                alignment=TA_CENTER
            )

            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=16,
                spaceAfter=20,
                spaceBefore=20
            )

            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=12
            )

            # Build content
            story = []

            # Title
            title = transcript.title or 'Transcript'
            story.append(Paragraph(title, title_style))

            # Metadata section
            if options.get('include_metadata', True):
                story.append(Paragraph('Case Information', heading_style))

                metadata_data = [
                    ['Language', transcript.language.upper()],
                    ['Duration', f"{transcript.totalDurationMs / 1000 / 60:.1f} minutes"],
                    ['Model', transcript.asrModel],
                    ['Created', transcript.createdAt],
                ]

                if matter:
                    metadata_data.extend([
                        ['Case', matter.title],
                        ['Client', matter.clientName or 'N/A'],
                        ['Practice Area', matter.practiceArea or 'N/A'],
                    ])

                metadata_table = Table(metadata_data, colWidths=[2*inch, 4*inch])
                metadata_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))

                story.append(metadata_table)
                story.append(Spacer(1, 20))

            # Transcript content
            story.append(Paragraph('Transcript', heading_style))

            for segment in transcript.segments:
                # Speaker and timestamp
                speaker_text = f"{segment.speaker} ({self._format_readable_time(segment.start_ms)})"
                story.append(Paragraph(speaker_text, styles['Heading3']))

                # Text content
                story.append(Paragraph(segment.text, normal_style))

                # Confidence score if enabled
                if options.get('include_confidence', True) and segment.confidence:
                    confidence_text = f"Confidence: {segment.confidence:.2%}"
                    story.append(Paragraph(confidence_text, styles['Italic']))

                story.append(Spacer(1, 12))

            # Build PDF
            doc.build(story)

            logger.info(f"PDF export completed: {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"PDF export failed: {e}")
            raise

    async def _export_csv(self, transcript: Transcript, options: Dict[str, Any]) -> str:
        """Export transcript in CSV format."""
        try:
            output_path = f"/tmp/transcript_{transcript.id}.csv"

            with open(output_path, 'w', encoding='utf-8', newline='') as f:
                # CSV header
                headers = ['Start Time', 'End Time', 'Speaker', 'Text']
                if options.get('include_confidence', True):
                    headers.append('Confidence')

                f.write(','.join(f'"{h}"' for h in headers) + '\n')

                # Data rows
                for segment in transcript.segments:
                    row = [
                        self._format_readable_time(segment.start_ms),
                        self._format_readable_time(segment.end_ms),
                        segment.speaker,
                        segment.text.replace('"', '""')  # Escape quotes
                    ]

                    if options.get('include_confidence', True):
                        row.append(f"{segment.confidence:.3f}" if segment.confidence else '')

                    f.write(','.join(f'"{field}"' for field in row) + '\n')

            logger.info(f"CSV export completed: {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"CSV export failed: {e}")
            raise

    async def _export_json(self, transcript: Transcript, options: Dict[str, Any]) -> str:
        """Export transcript in JSON format."""
        try:
            output_path = f"/tmp/transcript_{transcript.id}.json"

            # Prepare export data
            export_data = {
                'id': transcript.id,
                'title': transcript.title,
                'language': transcript.language,
                'asrModel': transcript.asrModel,
                'diarizationModel': transcript.diarizationModel,
                'totalDurationMs': transcript.totalDurationMs,
                'version': transcript.version,
                'createdAt': transcript.createdAt,
                'updatedAt': transcript.updatedAt,
                'segments': []
            }

            # Add segments
            for segment in transcript.segments:
                segment_data = {
                    'id': segment.id,
                    'speaker': segment.speaker,
                    'startMs': segment.startMs,
                    'endMs': segment.endMs,
                    'text': segment.text,
                    'confidence': segment.confidence
                }

                if options.get('include_words', True) and segment.words:
                    segment_data['words'] = segment.words

                export_data['segments'].append(segment_data)

            # Add speaker mapping if enabled
            if options.get('include_speaker_mapping', True) and transcript.speakerMap:
                export_data['speakerMap'] = transcript.speakerMap

            # Add metadata if enabled
            if options.get('include_metadata', True):
                export_data['metadata'] = {
                    'exportFormat': 'json',
                    'exportedAt': datetime.now().isoformat(),
                    'exportOptions': options
                }

            # Write JSON file
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(export_data, f, indent=2, ensure_ascii=False)

            logger.info(f"JSON export completed: {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"JSON export failed: {e}")
            raise

    def _format_srt_time(self, ms: int) -> str:
        """Format milliseconds to SRT timestamp format (HH:MM:SS,mmm)."""
        seconds = ms // 1000
        milliseconds = ms % 1000

        hours = seconds // 3600
        minutes = (seconds % 3600) // 60
        seconds = seconds % 60

        return f"{hours:02d}:{minutes:02d}:{seconds:02d},{milliseconds:03d}"

    def _format_vtt_time(self, ms: int) -> str:
        """Format milliseconds to VTT timestamp format (HH:MM:SS.mmm)."""
        seconds = ms // 1000
        milliseconds = ms % 1000

        hours = seconds // 3600
        minutes = (seconds % 3600) // 60
        seconds = seconds % 60

        return f"{hours:02d}:{minutes:02d}:{seconds:02d}.{milliseconds:03d}"

    def _format_readable_time(self, ms: int) -> str:
        """Format milliseconds to human-readable time format."""
        seconds = ms // 1000
        minutes = seconds // 60
        seconds = seconds % 60

        if minutes >= 60:
            hours = minutes // 60
            minutes = minutes % 60
            return f"{hours:02d}:{minutes:02d}:{seconds:02d}"
        else:
            return f"{minutes:02d}:{seconds:02d}"

    def get_supported_formats(self) -> List[str]:
        """Get list of supported export formats."""
        return self.supported_formats.copy()

    def get_format_info(self, format: str) -> Dict[str, Any]:
        """Get information about a specific export format."""
        format_info = {
            'srt': {
                'name': 'SRT (SubRip)',
                'description': 'Standard subtitle format for video players',
                'extension': '.srt',
                'options': ['include_speakers', 'include_timestamps']
            },
            'vtt': {
                'name': 'WebVTT',
                'description': 'Web video text tracks format',
                'extension': '.vtt',
                'options': ['include_speakers', 'include_timestamps']
            },
            'docx': {
                'name': 'Microsoft Word',
                'description': 'Editable document with formatting',
                'extension': '.docx',
                'options': ['include_metadata', 'include_speakers', 'include_confidence']
            },
            'pdf': {
                'name': 'PDF Document',
                'description': 'Professional quote pack format',
                'extension': '.pdf',
                'options': ['include_metadata', 'include_speakers', 'include_confidence']
            },
            'csv': {
                'name': 'CSV Spreadsheet',
                'description': 'Data analysis and processing',
                'extension': '.csv',
                'options': ['include_speakers', 'include_confidence']
            },
            'json': {
                'name': 'JSON Data',
                'description': 'Structured data with timestamps',
                'extension': '.json',
                'options': ['include_metadata', 'include_speaker_mapping', 'include_words']
            }
        }

        return format_info.get(format, {})
