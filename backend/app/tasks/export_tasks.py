"""
Background tasks for generating and exporting various file formats.
"""

import os
import tempfile
import asyncio
from typing import Dict, Any, List
from datetime import datetime
from pathlib import Path
from celery import current_task
from sqlalchemy.ext.asyncio import AsyncSession, create_async_engine
from sqlalchemy.orm import sessionmaker
from sqlalchemy import select

from app.tasks.celery_app import celery_app
from app.core.config import settings
from app.models.transcript import Transcript, TranscriptSegment
from app.models.matter import Matter
from app.models.user import User, Tenant
from app.services.storage_service import get_storage_service


# Create async database session for tasks
async_engine = create_async_engine(str(settings.DATABASE_URL))
AsyncSessionLocal = sessionmaker(
    async_engine, 
    class_=AsyncSession, 
    expire_on_commit=False
)


class ExportService:
    """Service for handling various export formats."""
    
    def __init__(self):
        self.storage_service = get_storage_service()
    
    def generate_plain_text(self, transcript: Transcript, segments: List[TranscriptSegment]) -> str:
        """Generate plain text format."""
        lines = []
        lines.append(f"Transcript: {transcript.title}")
        lines.append(f"Created: {transcript.created_at}")
        lines.append(f"Language: {transcript.language}")
        lines.append(f"Duration: {transcript.duration_minutes:.2f} minutes")
        lines.append("=" * 50)
        lines.append("")
        
        if transcript.has_speakers:
            current_speaker = None
            for segment in segments:
                if segment.speaker_name != current_speaker:
                    current_speaker = segment.speaker_name
                    lines.append(f"\n{segment.display_speaker}:")
                lines.append(f"{segment.text}")
        else:
            for segment in segments:
                lines.append(segment.text)
        
        return "\n".join(lines)
    
    def generate_timestamped_text(self, transcript: Transcript, segments: List[TranscriptSegment]) -> str:
        """Generate timestamped text format."""
        lines = []
        lines.append(f"Transcript: {transcript.title}")
        lines.append(f"Created: {transcript.created_at}")
        lines.append("=" * 50)
        lines.append("")
        
        for segment in segments:
            timestamp = f"[{segment.start_time_formatted}]"
            speaker = f"{segment.display_speaker}: " if transcript.has_speakers else ""
            lines.append(f"{timestamp} {speaker}{segment.text}")
        
        return "\n".join(lines)
    
    def generate_json_format(self, transcript: Transcript, segments: List[TranscriptSegment]) -> Dict[str, Any]:
        """Generate JSON format."""
        return {
            "transcript": {
                "id": str(transcript.id),
                "title": transcript.title,
                "language": transcript.language,
                "duration_seconds": transcript.duration_seconds,
                "word_count": transcript.word_count,
                "speaker_count": transcript.speaker_count,
                "created_at": transcript.created_at.isoformat(),
                "confidence_score": float(transcript.confidence_score) if transcript.confidence_score else None
            },
            "segments": [
                {
                    "index": segment.segment_index,
                    "start_time": segment.start_time,
                    "end_time": segment.end_time,
                    "duration": segment.duration,
                    "text": segment.text,
                    "speaker": {
                        "id": segment.speaker_id,
                        "name": segment.speaker_name,
                        "role": segment.speaker_role.value if segment.speaker_role else None
                    } if segment.has_speaker else None,
                    "confidence": float(segment.confidence) if segment.confidence else None,
                    "word_count": segment.word_count
                }
                for segment in segments
            ]
        }
    
    async def generate_pdf(self, transcript: Transcript, segments: List[TranscriptSegment]) -> bytes:
        """Generate PDF format using reportlab."""
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib.colors import black, blue
            import io
            
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=1*inch)
            
            # Get styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                textColor=blue,
                spaceAfter=20
            )
            
            speaker_style = ParagraphStyle(
                'Speaker',
                parent=styles['Normal'],
                fontSize=12,
                textColor=black,
                fontName='Helvetica-Bold',
                spaceBefore=10
            )
            
            # Build story
            story = []
            
            # Title and metadata
            story.append(Paragraph(transcript.title, title_style))
            story.append(Paragraph(f"Created: {transcript.created_at}", styles['Normal']))
            story.append(Paragraph(f"Language: {transcript.language}", styles['Normal']))
            story.append(Paragraph(f"Duration: {transcript.duration_minutes:.2f} minutes", styles['Normal']))
            story.append(Paragraph(f"Word Count: {transcript.word_count:,}", styles['Normal']))
            story.append(Spacer(1, 20))
            
            # Content
            if transcript.has_speakers:
                current_speaker = None
                for segment in segments:
                    if segment.speaker_name != current_speaker:
                        current_speaker = segment.speaker_name
                        story.append(Paragraph(f"{segment.display_speaker}:", speaker_style))
                    
                    text_with_timestamp = f"[{segment.start_time_formatted}] {segment.text}"
                    story.append(Paragraph(text_with_timestamp, styles['Normal']))
            else:
                for segment in segments:
                    text_with_timestamp = f"[{segment.start_time_formatted}] {segment.text}"
                    story.append(Paragraph(text_with_timestamp, styles['Normal']))
            
            # Build PDF
            doc.build(story)
            buffer.seek(0)
            
            return buffer.read()
            
        except ImportError:
            raise Exception("ReportLab not installed. Cannot generate PDF.")
        except Exception as e:
            raise Exception(f"PDF generation failed: {str(e)}")
    
    async def generate_docx(self, transcript: Transcript, segments: List[TranscriptSegment]) -> bytes:
        """Generate DOCX format using python-docx."""
        try:
            from docx import Document
            from docx.shared import Inches
            import io
            
            doc = Document()
            
            # Add title
            title = doc.add_heading(transcript.title, 0)
            
            # Add metadata
            doc.add_paragraph(f"Created: {transcript.created_at}")
            doc.add_paragraph(f"Language: {transcript.language}")
            doc.add_paragraph(f"Duration: {transcript.duration_minutes:.2f} minutes")
            doc.add_paragraph(f"Word Count: {transcript.word_count:,}")
            
            doc.add_page_break()
            
            # Add content
            if transcript.has_speakers:
                current_speaker = None
                for segment in segments:
                    if segment.speaker_name != current_speaker:
                        current_speaker = segment.speaker_name
                        speaker_para = doc.add_paragraph()
                        speaker_run = speaker_para.add_run(f"{segment.display_speaker}:")
                        speaker_run.bold = True
                    
                    text_with_timestamp = f"[{segment.start_time_formatted}] {segment.text}"
                    doc.add_paragraph(text_with_timestamp)
            else:
                for segment in segments:
                    text_with_timestamp = f"[{segment.start_time_formatted}] {segment.text}"
                    doc.add_paragraph(text_with_timestamp)
            
            # Save to bytes
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            return buffer.read()
            
        except ImportError:
            raise Exception("python-docx not installed. Cannot generate DOCX.")
        except Exception as e:
            raise Exception(f"DOCX generation failed: {str(e)}")


export_service = ExportService()


@celery_app.task(bind=True, name="app.tasks.export_tasks.export_transcript")
def export_transcript(self, transcript_id: str, format: str, options: Dict[str, Any] = None):
    """Export transcript in specified format."""
    return asyncio.run(_export_transcript_async(self, transcript_id, format, options or {}))


async def _export_transcript_async(task, transcript_id: str, format: str, options: Dict[str, Any]):
    """Async implementation of transcript export."""
    async with AsyncSessionLocal() as db:
        try:
            task.update_state(
                state="PROGRESS",
                meta={"stage": "initializing", "progress": 0}
            )
            
            # Get transcript with segments
            query = select(Transcript).where(Transcript.id == transcript_id)
            result = await db.execute(query)
            transcript = result.scalar_one_or_none()
            
            if not transcript:
                raise Exception(f"Transcript {transcript_id} not found")
            
            # Get segments
            segments_query = select(TranscriptSegment).where(
                TranscriptSegment.transcript_id == transcript_id
            ).order_by(TranscriptSegment.segment_index)
            
            segments_result = await db.execute(segments_query)
            segments = list(segments_result.scalars().all())
            
            task.update_state(
                state="PROGRESS",
                meta={"stage": "generating", "progress": 25}
            )
            
            # Generate export based on format
            export_content = None
            content_type = "text/plain"
            file_extension = ".txt"
            
            if format == "txt":
                if options.get("include_timestamps", True):
                    export_content = export_service.generate_timestamped_text(transcript, segments)
                else:
                    export_content = export_service.generate_plain_text(transcript, segments)
                export_content = export_content.encode('utf-8')
                content_type = "text/plain"
                file_extension = ".txt"
                
            elif format == "json":
                import json
                export_data = export_service.generate_json_format(transcript, segments)
                export_content = json.dumps(export_data, indent=2, ensure_ascii=False).encode('utf-8')
                content_type = "application/json"
                file_extension = ".json"
                
            elif format == "pdf":
                export_content = await export_service.generate_pdf(transcript, segments)
                content_type = "application/pdf"
                file_extension = ".pdf"
                
            elif format == "docx":
                export_content = await export_service.generate_docx(transcript, segments)
                content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                file_extension = ".docx"
                
            else:
                raise Exception(f"Unsupported export format: {format}")
            
            task.update_state(
                state="PROGRESS",
                meta={"stage": "storing", "progress": 75}
            )
            
            # Store export file
            timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
            filename = f"{transcript.title}_{timestamp}{file_extension}"
            
            # Store in exports directory
            export_path = f"exports/{transcript.tenant_id}/{filename}"
            storage_path = export_service.storage_service.storage_root / export_path
            storage_path.parent.mkdir(parents=True, exist_ok=True)
            
            with open(storage_path, "wb") as f:
                f.write(export_content)
            
            # Mark transcript as exported
            transcript.mark_exported(format)
            await db.commit()
            
            task.update_state(
                state="PROGRESS",
                meta={"stage": "completed", "progress": 100}
            )
            
            return {
                "success": True,
                "export_path": export_path,
                "filename": filename,
                "file_size": len(export_content),
                "content_type": content_type,
                "download_url": f"/api/v1/exports/download/{export_path}"
            }
            
        except Exception as e:
            raise e


@celery_app.task(name="app.tasks.export_tasks.export_matter_summary")
def export_matter_summary(matter_id: str, format: str = "pdf"):
    """Export matter summary with all transcripts."""
    return asyncio.run(_export_matter_summary_async(matter_id, format))


async def _export_matter_summary_async(matter_id: str, format: str):
    """Generate comprehensive matter summary export."""
    async with AsyncSessionLocal() as db:
        try:
            # Get matter with all related data
            query = select(Matter).where(Matter.id == matter_id)
            result = await db.execute(query)
            matter = result.scalar_one_or_none()
            
            if not matter:
                raise Exception(f"Matter {matter_id} not found")
            
            # Get all transcripts for this matter
            transcripts_query = select(Transcript).where(
                Transcript.matter_id == matter_id
            ).order_by(Transcript.created_at)
            
            transcripts_result = await db.execute(transcripts_query)
            transcripts = list(transcripts_result.scalars().all())
            
            # TODO: Generate comprehensive matter summary
            # This would include:
            # - Matter details
            # - All transcripts
            # - Summary statistics
            # - Timeline of events
            
            return {
                "success": True,
                "matter_title": matter.title,
                "transcript_count": len(transcripts),
                "format": format
            }
            
        except Exception as e:
            raise e


@celery_app.task(name="app.tasks.export_tasks.cleanup_old_exports")
def cleanup_old_exports(max_age_days: int = 7):
    """Clean up old export files."""
    storage_service = get_storage_service()
    exports_path = storage_service.storage_root / "exports"
    
    if not exports_path.exists():
        return {"deleted_files": 0}
    
    import time
    cutoff_time = time.time() - (max_age_days * 24 * 3600)
    deleted_count = 0
    
    for file_path in exports_path.rglob("*"):
        if file_path.is_file() and file_path.stat().st_mtime < cutoff_time:
            try:
                file_path.unlink()
                deleted_count += 1
            except OSError:
                pass
    
    return {"deleted_files": deleted_count, "max_age_days": max_age_days}